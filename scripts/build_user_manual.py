"""Build TaxFlow Pro user manual as PDF and Word document.

Usage:
    python scripts/build_user_manual.py

Outputs:
    projects/TaxFlow-Pro/TaxFlow-Pro-v3.9/docs/TaxFlow-Pro-v3.11.6-User-Manual.pdf
    projects/TaxFlow-Pro/TaxFlow-Pro-v3.9/docs/TaxFlow-Pro-v3.11.6-User-Manual.docx
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    PageBreak,
    ListFlowable,
    ListItem,
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
SOURCE_MD = DOCS_DIR / "user-manual.md"
PDF_OUT = DOCS_DIR / "TaxFlow-Pro-v3.11.6-User-Manual.pdf"
DOCX_OUT = DOCS_DIR / "TaxFlow-Pro-v3.11.6-User-Manual.docx"


def parse_markdown(path: Path) -> list[dict]:
    """Parse a subset of Markdown into blocks suitable for both PDF and DOCX."""
    text = path.read_text(encoding="utf-8")
    blocks: list[dict] = []

    # Pre-process: protect code blocks
    code_blocks: list[str] = []

    def stash_code(match: re.Match) -> str:
        code_blocks.append(match.group(1))
        return f"\n\n__CODE_BLOCK_{len(code_blocks) - 1}__\n\n"

    text = re.sub(r"```[a-zA-Z0-9_+-]*\n(.*?)\n```", stash_code, text, flags=re.S)

    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("__CODE_BLOCK_"):
            idx = int(stripped.replace("__CODE_BLOCK_", "").replace("__", ""))
            blocks.append({"type": "code", "text": code_blocks[idx]})
            i += 1
            continue

        # Header
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            title = stripped.lstrip("#").strip()
            blocks.append({"type": f"h{level}", "text": title})
            i += 1
            continue

        # Horizontal rule
        if stripped == "---":
            blocks.append({"type": "hr"})
            i += 1
            continue

        # Blockquote / callout
        if stripped.startswith(">"):
            callout_lines = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                callout_lines.append(lines[i].strip().lstrip(">").strip())
                i += 1
            blocks.append({"type": "callout", "text": " ".join(callout_lines)})
            continue

        # Table (not implemented in this converter; treat as preformatted)
        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            blocks.append({"type": "pre", "text": "\n".join(table_lines)})
            continue

        # List item
        if re.match(r"^(\s*)([-*]|\d+\.)\s+", stripped):
            list_items: list[str] = []
            list_type = "ul"
            while i < len(lines):
                s = lines[i].strip()
                if not s:
                    i += 1
                    continue
                m = re.match(r"^\s*([-*]|\d+\.)\s+(.*)$", s)
                if not m:
                    break
                if re.match(r"^\d+\.$", m.group(1)):
                    list_type = "ol"
                list_items.append(m.group(2))
                i += 1
            blocks.append({"type": list_type, "items": list_items})
            continue

        # Regular paragraph (possibly with inline formatting)
        para_lines = []
        while i < len(lines) and lines[i].strip():
            para_lines.append(lines[i])
            i += 1
        para = " ".join(para_lines)
        blocks.append({"type": "p", "text": para})

    return blocks


def _inline_to_reportlab(text: str) -> str:
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
    text = re.sub(r"`([^`]+?)`", r"<font face='Courier'>\1</font>", text)
    text = re.sub(r"\*(.+?)\*", r"<i>\1</i>", text)
    return text


def _inline_to_docx(text: str) -> list[tuple[str, str]]:
    """Split text into (style, fragment) tuples."""
    result: list[tuple[str, str]] = []
    pos = 0
    patterns = [
        (r"\*\*(.+?)\*\*", "bold"),
        (r"`([^`]+?)`", "code"),
        (r"\*(.+?)\*", "italic"),
    ]
    while pos < len(text):
        best: tuple[int, str, str, re.Match | None] = (-1, "", "", None)
        for pattern, style in patterns:
            m = re.search(pattern, text[pos:])
            if m:
                start = pos + m.start()
                if best[0] == -1 or start < best[0]:
                    best = (start, style, m.group(1), m)
        if best[0] == -1:
            result.append(("", text[pos:]))
            break
        if best[0] > pos:
            result.append(("", text[pos:best[0]]))
        result.append((best[1], best[2]))
        pos = best[0] + len(best[3].group(0))
    return result


def _slug(text: str) -> str:
    return re.sub(r"[^a-z0-9]+", "-", text.lower()).strip("-")


def build_pdf(blocks: list[dict], out_path: Path) -> None:
    doc = SimpleDocTemplate(
        str(out_path),
        pagesize=letter,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )

    styles = getSampleStyleSheet()
    if "ManualTitle" not in styles:
        styles.add(
            ParagraphStyle(
                "ManualTitle",
                parent=styles["Heading1"],
                fontSize=24,
                leading=30,
                alignment=TA_CENTER,
                spaceAfter=18,
            )
        )
    styles.add(
        ParagraphStyle(
            "Subtitle",
            parent=styles["Normal"],
            fontSize=12,
            leading=16,
            alignment=TA_CENTER,
            textColor=RGBColor(0x55, 0x55, 0x55),
            spaceAfter=24,
        )
    )
    styles.add(
        ParagraphStyle(
            "BodyJustify",
            parent=styles["Normal"],
            alignment=TA_JUSTIFY,
            leading=14,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            "Callout",
            parent=styles["Normal"],
            leftIndent=12,
            rightIndent=12,
            backColor=RGBColor(0xF5, 0xF5, 0xF5),
            borderPadding=8,
            spaceAfter=10,
            leading=14,
        )
    )
    if "CodeBlock" not in styles:
        styles.add(
            ParagraphStyle(
                "CodeBlock",
                parent=styles["Code"],
                fontSize=9,
                leading=11,
                leftIndent=12,
                rightIndent=12,
                backColor=RGBColor(0xF0, 0xF0, 0xF0),
                borderPadding=6,
                spaceAfter=10,
            )
        )

    story: list = []

    # Cover page
    story.append(Paragraph("TaxFlow Pro", styles["ManualTitle"]))
    story.append(Paragraph("v3.11.6 — User Manual", styles["ManualTitle"]))
    story.append(Spacer(1, 0.25 * inch))
    story.append(
        Paragraph(
            "Local-first, offline-capable financial document processing for individuals and small businesses.",
            styles["Subtitle"],
        )
    )
    story.append(Spacer(1, 0.5 * inch))
    story.append(
        Paragraph(
            "Last updated: 2026-07-01",
            styles["Subtitle"],
        )
    )
    story.append(PageBreak())

    # Table of contents (generated from h1/h2 blocks)
    story.append(Paragraph("Table of Contents", styles["Heading1"]))
    for block in blocks:
        if block["type"] in ("h1", "h2"):
            indent = 0 if block["type"] == "h1" else 18
            story.append(
                Paragraph(
                    block["text"],
                    ParagraphStyle(
                        f"TOC{block['type']}",
                        parent=styles["Normal"],
                        leftIndent=indent,
                        spaceAfter=4,
                        fontSize=12 if block["type"] == "h1" else 11,
                    ),
                )
            )
    story.append(PageBreak())

    for block in blocks:
        t = block["type"]
        if t == "h1":
            story.append(Paragraph(block["text"], styles["Heading1"]))
        elif t == "h2":
            story.append(Paragraph(block["text"], styles["Heading2"]))
        elif t == "h3":
            story.append(Paragraph(block["text"], styles["Heading3"]))
        elif t == "p":
            story.append(
                Paragraph(_inline_to_reportlab(block["text"]), styles["BodyJustify"])
            )
        elif t in ("ul", "ol"):
            items = [
                ListItem(Paragraph(_inline_to_reportlab(item), styles["BodyJustify"]))
                for item in block["items"]
            ]
            story.append(
                ListFlowable(
                    items,
                    bulletType="bullet" if t == "ul" else "1",
                    leftIndent=18,
                    bulletFontSize=10,
                    bulletOffsetY=-2,
                )
            )
        elif t == "code":
            for line in block["text"].splitlines():
                story.append(
                    Paragraph(
                        line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"),
                        styles["CodeBlock"],
                    )
                )
        elif t == "callout":
            story.append(
                Paragraph(
                    f"<b>Tip:</b> {_inline_to_reportlab(block['text'])}",
                    styles["Callout"],
                )
            )
        elif t == "hr":
            story.append(Spacer(1, 0.2 * inch))

    doc.build(story)


def build_docx(blocks: list[dict], out_path: Path) -> None:
    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Title
    title = doc.add_heading("TaxFlow Pro", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("v3.11.6 — User Manual")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph(
        "Local-first, offline-capable financial document processing for individuals and small businesses."
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Last updated: 2026-07-01").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # Table of contents
    doc.add_heading("Table of Contents", level=1)
    for block in blocks:
        if block["type"] in ("h1", "h2"):
            p = doc.add_paragraph()
            if block["type"] == "h2":
                p.paragraph_format.left_indent = Inches(0.25)
            run = p.add_run(block["text"])
            run.font.size = Pt(13 if block["type"] == "h1" else 11)
            if block["type"] == "h1":
                run.bold = True
    doc.add_page_break()

    for block in blocks:
        t = block["type"]
        if t in ("h1", "h2", "h3"):
            level = {"h1": 1, "h2": 2, "h3": 3}[t]
            doc.add_heading(block["text"], level=level)
        elif t == "p":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            for style_name, fragment in _inline_to_docx(block["text"]):
                run = p.add_run(fragment)
                if style_name == "bold":
                    run.bold = True
                elif style_name == "italic":
                    run.italic = True
                elif style_name == "code":
                    run.font.name = "Courier New"
                    run.font.size = Pt(10)
        elif t in ("ul", "ol"):
            for idx, item in enumerate(block["items"], start=1):
                p = doc.add_paragraph(style="List Bullet" if t == "ul" else "List Number")
                p.paragraph_format.space_after = Pt(4)
                for style_name, fragment in _inline_to_docx(item):
                    run = p.add_run(fragment)
                    if style_name == "bold":
                        run.bold = True
                    elif style_name == "italic":
                        run.italic = True
                    elif style_name == "code":
                        run.font.name = "Courier New"
                        run.font.size = Pt(10)
        elif t == "code":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(block["text"])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        elif t == "pre":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(block["text"])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        elif t == "callout":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.right_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(8)
            r = p.add_run("Tip: ")
            r.bold = True
            p.add_run(block["text"])
        elif t == "hr":
            doc.add_paragraph()

    doc.save(str(out_path))


def main() -> int:
    if not SOURCE_MD.exists():
        print(f"Source Markdown not found: {SOURCE_MD}", file=sys.stderr)
        return 1

    blocks = parse_markdown(SOURCE_MD)
    build_pdf(blocks, PDF_OUT)
    build_docx(blocks, DOCX_OUT)
    print(f"Created: {PDF_OUT}")
    print(f"Created: {DOCX_OUT}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
